"""Servizio per l'export del libro in diversi formati (EPUB, DOCX)."""
from pathlib import Path
from io import BytesIO
from datetime import datetime
import base64
import markdown
from typing import Optional
from PIL import Image as PILImage
import ebooklib
from ebooklib import epub
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.agent.session_store import SessionData
from app.services.pdf_service import (
    get_model_abbreviation,
    escape_html,
    markdown_to_html,
)
from app.services.storage_service import get_storage_service


def generate_epub(session: SessionData) -> tuple[bytes, str]:
    """
    Genera un file EPUB del libro completo con titolo, indice e capitoli.
    
    Args:
        session: SessionData object
        
    Returns:
        Tupla (epub_bytes, filename)
    """
    # Prepara dati libro
    book_title = session.current_title or "Romanzo"
    book_author = session.form_data.user_name or "Autore"
    
    # Crea il libro EPUB
    book = epub.EpubBook()
    
    # Metadati
    book.set_identifier(session.session_id)
    book.set_title(book_title)
    book.set_language('it')
    book.add_author(book_author)
    
    # Aggiungi copertina se disponibile
    cover_image_data = None
    cover_mime_type = 'image/png'
    if session.cover_image_path:
        try:
            storage_service = get_storage_service()
            print(f"[EPUB] Caricamento copertina da: {session.cover_image_path}")
            cover_image_data = storage_service.download_file(session.cover_image_path)
            
            print(f"[EPUB] Immagine copertina caricata: {len(cover_image_data)} bytes")
            
            # Determina MIME type e estensione dal path
            cover_path_str = session.cover_image_path
            cover_file_extension = '.jpg'
            if '.png' in cover_path_str.lower():
                cover_mime_type = 'image/png'
                cover_file_extension = '.png'
            elif '.jpg' in cover_path_str.lower() or '.jpeg' in cover_path_str.lower():
                cover_mime_type = 'image/jpeg'
                cover_file_extension = '.jpg'
            else:
                cover_mime_type = 'image/png'
                cover_file_extension = '.png'
            
            print(f"[EPUB] Tipo MIME copertina: {cover_mime_type}")
            
            # Imposta copertina nei metadati
            book.set_cover(f'cover{cover_file_extension}', cover_image_data)
            print(f"[EPUB] Copertina impostata nei metadati EPUB")
        except Exception as e:
            print(f"[EPUB] Errore nel caricamento copertina: {e}")
            import traceback
            traceback.print_exc()
    else:
        print(f"[EPUB] Nessuna copertina disponibile (path: {session.cover_image_path})")
    
    # Ordina i capitoli per section_index
    sorted_chapters = sorted(session.book_chapters, key=lambda x: x.get('section_index', 0))
    
    # Crea capitolo copertina come prima pagina se disponibile
    cover_chapter = None
    if cover_image_data:
        try:
            print(f"[EPUB] Creazione capitolo copertina...")
            # Usa base64 per includere l'immagine direttamente nell'HTML
            cover_base64 = base64.b64encode(cover_image_data).decode('utf-8')
            print(f"[EPUB] Immagine convertita in base64: {len(cover_base64)} caratteri")
            
            cover_html = f'''<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>Copertina</title>
    <meta charset="UTF-8"/>
    <style>
        @page {{
            margin: 0;
            padding: 0;
        }}
        body {{
            margin: 0;
            padding: 0;
            display: flex;
            align-items: center;
            justify-content: center;
            min-height: 100vh;
            width: 100%;
            background-color: #000;
            overflow: hidden;
        }}
        img {{
            max-width: 100%;
            max-height: 100vh;
            width: auto;
            height: auto;
            display: block;
            object-fit: contain;
        }}
    </style>
</head>
<body>
    <img src="data:{cover_mime_type};base64,{cover_base64}" alt="Copertina" />
</body>
</html>'''
            
            cover_chapter = epub.EpubHtml(
                title='Copertina',
                file_name='cover.xhtml',
                lang='it'
            )
            cover_chapter.content = cover_html
            book.add_item(cover_chapter)
            print(f"[EPUB] Capitolo copertina creato e aggiunto al libro")
        except Exception as e:
            print(f"[EPUB] Errore nella creazione del capitolo copertina: {e}")
            import traceback
            traceback.print_exc()
    else:
        print(f"[EPUB] Nessuna immagine copertina disponibile, salto creazione capitolo")
    
    # Crea capitoli EPUB
    epub_chapters = []
    toc_items = []
    
    for idx, chapter in enumerate(sorted_chapters, 1):
        chapter_title = chapter.get('title', f'Capitolo {idx}')
        chapter_content = chapter.get('content', '')
        
        # Converti markdown a HTML per EPUB
        content_html = markdown_to_html(chapter_content)
        
        # Crea il capitolo HTML completo
        chapter_html = f'''<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>{escape_html(chapter_title)}</title>
    <style>
        body {{ font-family: serif; margin: 1em; line-height: 1.6; }}
        h1 {{ font-size: 1.8em; margin-bottom: 0.5em; border-bottom: 1px solid #ccc; padding-bottom: 0.3em; }}
        p {{ margin: 0.8em 0; text-align: justify; }}
    </style>
</head>
<body>
    <h1>{escape_html(chapter_title)}</h1>
    <div class="chapter-content">
        {content_html}
    </div>
</body>
</html>'''
        
        # Crea il capitolo EPUB
        chapter_file = epub.EpubHtml(
            title=chapter_title,
            file_name=f'chapter_{idx}.xhtml',
            lang='it'
        )
        chapter_file.content = chapter_html
        
        book.add_item(chapter_file)
        epub_chapters.append(chapter_file)
        toc_items.append(chapter_file)
    
    # Aggiungi indice dei contenuti (TOC)
    book.toc = tuple([(epub.Section('Indice'), tuple(toc_items))])
    
    # Aggiungi spine (ordine di lettura) - copertina come prima pagina se disponibile
    spine_items = ['nav']
    if cover_chapter:
        spine_items.append(cover_chapter)
        print(f"[EPUB] Copertina aggiunta allo spine come prima pagina")
    spine_items.extend(epub_chapters)
    book.spine = spine_items
    print(f"[EPUB] Spine configurato con {len(spine_items)} elementi totali (nav + {'copertina + ' if cover_chapter else ''}{len(epub_chapters)} capitoli)")
    
    # Aggiungi navigazione
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    
    # Genera EPUB in memoria
    buffer = BytesIO()
    epub.write_epub(buffer, book, {})
    buffer.seek(0)
    epub_bytes = buffer.getvalue()
    
    # Nome file
    date_prefix = datetime.now().strftime("%Y-%m-%d")
    model_abbrev = get_model_abbreviation(session.form_data.llm_model)
    title_sanitized = "".join(c for c in book_title if c.isalnum() or c in (' ', '-', '_')).rstrip()
    title_sanitized = title_sanitized.replace(" ", "_")
    if not title_sanitized:
        title_sanitized = f"Libro_{session.session_id[:8]}"
    filename = f"{date_prefix}_{model_abbrev}_{title_sanitized}.epub"
    
    return epub_bytes, filename


def generate_docx(session: SessionData) -> tuple[bytes, str]:
    """
    Genera un file DOCX del libro completo con titolo, indice e capitoli.
    
    Args:
        session: SessionData object
        
    Returns:
        Tupla (docx_bytes, filename)
    """
    # Prepara dati libro
    book_title = session.current_title or "Romanzo"
    book_author = session.form_data.user_name or "Autore"
    
    # Crea il documento
    doc = Document()
    
    # Configura stili predefiniti
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Titolo del libro
    title_para = doc.add_paragraph(book_title)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Autore
    author_para = doc.add_paragraph(book_author)
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.runs[0]
    author_run.font.size = Pt(14)
    author_run.font.italic = True
    
    # Aggiungi immagine copertina se disponibile
    if session.cover_image_path:
        try:
            storage_service = get_storage_service()
            cover_image_data = storage_service.download_file(session.cover_image_path)
            doc.add_page_break()
            cover_para = doc.add_paragraph()
            cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cover_para.add_run()
            # Salva temporaneamente in BytesIO per add_picture
            from io import BytesIO
            cover_io = BytesIO(cover_image_data)
            run.add_picture(cover_io, width=Inches(5))
            cover_io.close()
            doc.add_page_break()
        except Exception as e:
            print(f"[DOCX] Errore nel caricamento copertina: {e}")
    
    # Aggiungi indice
    doc.add_paragraph("Indice").runs[0].font.bold = True
    doc.add_paragraph()  # Riga vuota dopo il titolo
    
    # Ordina i capitoli per section_index
    sorted_chapters = sorted(session.book_chapters, key=lambda x: x.get('section_index', 0))
    
    # Aggiungi voci all'indice
    for idx, chapter in enumerate(sorted_chapters, 1):
        chapter_title = chapter.get('title', f'Capitolo {idx}')
        toc_para = doc.add_paragraph(f"{idx}. {chapter_title}", style='List Number')
        toc_para.runs[0].font.size = Pt(12)
    
    doc.add_page_break()
    
    # Aggiungi capitoli
    for idx, chapter in enumerate(sorted_chapters, 1):
        chapter_title = chapter.get('title', f'Capitolo {idx}')
        chapter_content = chapter.get('content', '')
        
        # Titolo del capitolo
        chapter_heading = doc.add_heading(chapter_title, level=1)
        chapter_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Converti markdown a testo semplice per DOCX
        # Rimuovi formattazione markdown base
        content_text = chapter_content
        # Rimuovi header markdown
        lines = content_text.split('\n')
        cleaned_lines = []
        for line in lines:
            # Rimuovi # headers
            while line.startswith('#'):
                line = line[1:].lstrip()
            # Rimuovi ** bold e * italic
            line = line.replace('**', '').replace('*', '').replace('__', '').replace('_', '')
            cleaned_lines.append(line)
        content_text = '\n'.join(cleaned_lines)
        
        # Aggiungi contenuto come paragrafi
        paragraphs = content_text.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                # Gestisci liste semplici
                if para_text.startswith('- ') or para_text.startswith('* '):
                    list_item = para_text[2:].strip()
                    doc.add_paragraph(list_item, style='List Bullet')
                elif para_text.startswith(tuple([f'{i}. ' for i in range(1, 10)])):
                    list_item = para_text.split('. ', 1)[1] if '. ' in para_text else para_text
                    doc.add_paragraph(list_item, style='List Number')
                else:
                    para = doc.add_paragraph(para_text)
                    para.runs[0].font.size = Pt(12)
                    para.paragraph_format.line_spacing = 1.15
        
        # Aggiungi spazio tra capitoli
        if idx < len(sorted_chapters):
            doc.add_page_break()
    
    # Salva in memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    docx_bytes = buffer.getvalue()
    
    # Nome file
    date_prefix = datetime.now().strftime("%Y-%m-%d")
    model_abbrev = get_model_abbreviation(session.form_data.llm_model)
    title_sanitized = "".join(c for c in book_title if c.isalnum() or c in (' ', '-', '_')).rstrip()
    title_sanitized = title_sanitized.replace(" ", "_")
    if not title_sanitized:
        title_sanitized = f"Libro_{session.session_id[:8]}"
    filename = f"{date_prefix}_{model_abbrev}_{title_sanitized}.docx"
    
    return docx_bytes, filename
